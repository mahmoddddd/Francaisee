"""Professional two-column Arabic CV for Mohamed Ali Salem — Land Surveyor."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Palette ----------
NAVY = RGBColor(0x0A, 0x2A, 0x4A)       # Deep navy — main dark
ACCENT = RGBColor(0xC9, 0xA0, 0x5E)     # Warm gold accent (very executive)
DARK = RGBColor(0x1E, 0x1E, 0x1E)
GRAY = RGBColor(0x5A, 0x5A, 0x5A)
LIGHT_GRAY = RGBColor(0x8C, 0x8C, 0x8C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SIDEBAR_TEXT = RGBColor(0xE8, 0xEE, 0xF5)

AR_FONT = 'Arial'


# ---------- XML helpers ----------
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_margins(cell, top=0.3, bottom=0.3, left=0.5, right=0.5):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val_cm in (('top', top), ('bottom', bottom),
                         ('left', left), ('right', right)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(int(val_cm * 567)))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tc_pr.append(tcMar)


def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def make_table_rtl(table):
    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    bidi.set(qn('w:val'), '1')
    tblPr.append(bidi)


def set_rtl(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    p_pr.append(bidi)


def set_run_rtl(run):
    r_pr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    r_pr.append(rtl)


def set_run_complex_font(run, font_name, size_pt):
    r_pr = run._r.get_or_add_rPr()
    rFonts = r_pr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r_pr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    r_pr.append(szCs)


def ar_run(paragraph, text, size=10.5, bold=False, color=DARK, font=AR_FONT,
           italic=False):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    set_run_complex_font(run, font, size)
    set_run_rtl(run)
    return run


def add_border(paragraph, position='bottom', color_hex='C9A05E', size='8'):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = paragraph._p.find(qn('w:pBdr'))
    if p_bdr is None:
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)
    border = OxmlElement(f'w:{position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), size)
    border.set(qn('w:space'), '2')
    border.set(qn('w:color'), color_hex)
    p_bdr.append(border)


# ---------- Sidebar helpers (dark navy side) ----------
def sidebar_section_title(cell, text, first=False):
    p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0 if first else 12)
    p.paragraph_format.space_after = Pt(3)
    ar_run(p, text, size=12, bold=True, color=ACCENT)
    add_border(p, 'bottom', color_hex='C9A05E', size='6')
    return p


def sidebar_line(cell, text, size=10, bold=False, color=SIDEBAR_TEXT,
                 space_after=2, align='right'):
    p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if align == 'right'
                   else WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    ar_run(p, text, size=size, bold=bold, color=color)
    return p


def sidebar_bullet(cell, text, size=10):
    p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    marker = p.add_run('■  ')
    marker.font.name = AR_FONT
    marker.font.size = Pt(7)
    marker.font.color.rgb = ACCENT
    marker.bold = True
    set_run_complex_font(marker, AR_FONT, 7)
    set_run_rtl(marker)
    ar_run(p, text, size=size, color=SIDEBAR_TEXT)


# ---------- Main area helpers (light side) ----------
def main_section_title(cell, text, first=False):
    p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0 if first else 10)
    p.paragraph_format.space_after = Pt(3)
    ar_run(p, text, size=14, bold=True, color=NAVY)
    add_border(p, 'bottom', color_hex='0A2A4A', size='6')
    return p


def main_paragraph(cell, text, size=10.5, color=DARK, bold=False,
                   italic=False, align='justify', space_after=3,
                   space_before=0):
    p = cell.add_paragraph()
    set_rtl(p)
    alignment_map = {
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
    }
    p.alignment = alignment_map[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    ar_run(p, text, size=size, bold=bold, color=color, italic=italic)
    return p


def main_bullet(cell, text, size=10.5):
    p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    marker = p.add_run('◆  ')
    marker.font.name = AR_FONT
    marker.font.size = Pt(8)
    marker.font.color.rgb = ACCENT
    marker.bold = True
    set_run_complex_font(marker, AR_FONT, 8)
    set_run_rtl(marker)
    ar_run(p, text, size=size, color=DARK)


# ---------- Build document ----------
doc = Document()

# Tight margins for two-column layout
for section in doc.sections:
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)

style = doc.styles['Normal']
style.font.name = AR_FONT
style.font.size = Pt(10.5)

# ========= Top Header (full width) =========
header_tbl = doc.add_table(rows=1, cols=1)
header_tbl.autofit = False
header_tbl.columns[0].width = Cm(21.0)
hc = header_tbl.cell(0, 0)
hc.width = Cm(21.0)
set_cell_bg(hc, '0A2A4A')
set_cell_margins(hc, top=0.6, bottom=0.6, left=1.2, right=1.2)

# Name
p_name = hc.paragraphs[0]
set_rtl(p_name)
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_name.paragraph_format.space_before = Pt(0)
p_name.paragraph_format.space_after = Pt(2)
ar_run(p_name, 'محمد علي سالم', size=30, bold=True, color=WHITE)

# Gold divider paragraph
p_div = hc.add_paragraph()
set_rtl(p_div)
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_div.paragraph_format.space_after = Pt(2)
divider_run = p_div.add_run('━━━━━━━━━━')
divider_run.font.size = Pt(10)
divider_run.font.color.rgb = ACCENT
divider_run.font.name = AR_FONT

# Title
p_title = hc.add_paragraph()
set_rtl(p_title)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(0)
ar_run(p_title, 'مهندس مساحة  -  مساح موقع', size=14,
       bold=True, color=ACCENT)

remove_table_borders(header_tbl)

# ========= Two-column body =========
body_tbl = doc.add_table(rows=1, cols=2)
body_tbl.autofit = False
# In default (non-RTL) table layout cell(0,0) is visually LEFT, cell(0,1) RIGHT
# We make the table bidiVisual so cell(0,0) becomes visually RIGHT.
# Right cell (RTL visual right) = sidebar with contact/skills
# Left cell  (RTL visual left)  = main content with experience

main_cell = body_tbl.cell(0, 0)       # Visually on LEFT side (after bidi)
side_cell = body_tbl.cell(0, 1)       # Wait — with bidiVisual this becomes the LEFT
# Let's instead set widths directly and rely on bidiVisual to mirror.

# Actually the cleanest: don't use bidiVisual on the table, just place
# sidebar in cell(0,1) = right side visually (since default LTR table).
# Arabic paragraphs inside will be RTL — the table layout itself LTR is fine.

# So: reassign for clarity.
main_cell = body_tbl.cell(0, 0)   # LEFT column visually — main content
side_cell = body_tbl.cell(0, 1)   # RIGHT column visually — sidebar

# Set column widths
body_tbl.columns[0].width = Cm(13.0)
body_tbl.columns[1].width = Cm(8.0)
main_cell.width = Cm(13.0)
side_cell.width = Cm(8.0)

# Sidebar styling
set_cell_bg(side_cell, '0A2A4A')
set_cell_margins(side_cell, top=0.7, bottom=0.7, left=0.7, right=0.7)

# Main cell styling
set_cell_margins(main_cell, top=0.7, bottom=0.7, left=0.8, right=0.8)

# Align cell contents to top
main_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
side_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

remove_table_borders(body_tbl)

# ======== Fill Sidebar (right) ========
# Clear default empty paragraph
side_cell.paragraphs[0].text = ''

sidebar_section_title(side_cell, 'بيانات التواصل', first=True)
sidebar_line(side_cell, '📞  الهاتف', size=9, bold=True, color=ACCENT,
             space_after=0)
sidebar_line(side_cell, '01032169010', size=10, color=SIDEBAR_TEXT,
             space_after=5)
sidebar_line(side_cell, '✉  البريد الإلكتروني', size=9, bold=True,
             color=ACCENT, space_after=0)
sidebar_line(side_cell, 'mimo99999911@gmail.com', size=10,
             color=SIDEBAR_TEXT, space_after=5)
sidebar_line(side_cell, '📍  العنوان', size=9, bold=True, color=ACCENT,
             space_after=0)
sidebar_line(side_cell, 'القاهرة، مصر', size=10, color=SIDEBAR_TEXT,
             space_after=3)

sidebar_section_title(side_cell, 'المهارات الفنية')
for skill in [
    'تشغيل جهاز التوتال استيشن (Total Station)',
    'تشغيل جهاز الميزان (Auto Level)',
    'تشغيل أجهزة الـ GPS / GNSS / RTK',
    'أعمال التوقيع والرفع المساحي',
    'الرفع الطبوغرافي',
    'أعمال الـ As-Built',
    'حساب الكميات والمناسيب',
    'قراءة الرسومات التنفيذية',
    'برنامج AutoCAD',
]:
    sidebar_bullet(side_cell, skill, size=10)

sidebar_section_title(side_cell, 'المؤهل العلمي')
sidebar_line(side_cell, 'دبلوم مساحة', size=11, bold=True,
             color=WHITE, space_after=1)
sidebar_line(side_cell, 'سنة التخرج: 2023', size=10,
             color=SIDEBAR_TEXT, space_after=3)

sidebar_section_title(side_cell, 'اللغات')
sidebar_line(side_cell, 'العربية', size=10, bold=True, color=WHITE,
             space_after=0)
sidebar_line(side_cell, 'اللغة الأم', size=9, color=SIDEBAR_TEXT,
             space_after=4)
sidebar_line(side_cell, 'الإنجليزية', size=10, bold=True, color=WHITE,
             space_after=0)
sidebar_line(side_cell, 'مستوى جيد (فني / تقني)', size=9,
             color=SIDEBAR_TEXT, space_after=3)

sidebar_section_title(side_cell, 'المهارات الشخصية')
for s in [
    'الدقة والاهتمام بالتفاصيل',
    'الالتزام والمسؤولية',
    'العمل الجماعي',
    'إدارة الوقت والأولويات',
    'التواصل الفعّال مع فرق العمل',
]:
    sidebar_bullet(side_cell, s, size=10)

# ======== Fill Main Area (left) ========
main_cell.paragraphs[0].text = ''

# ---- Summary ----
main_section_title(main_cell, 'الملخص المهني', first=True)
main_paragraph(
    main_cell,
    'فني مساحة يتمتع بخبرة عملية تزيد عن 3 سنوات في تنفيذ أعمال المساحة '
    'لمشروعات إنشائية وبنية تحتية كبرى داخل جمهورية مصر العربية. يجمع بين '
    'الدقة الفنية العالية والخبرة الميدانية الواسعة مع كبرى شركات المقاولات، '
    'ويجيد تشغيل التوتال استيشن والميزان وأجهزة الـ GPS / RTK. يعمل بكفاءة '
    'تحت ضغط الجداول الزمنية، ولديه قدرة عالية على قراءة الرسومات التنفيذية '
    'ونقلها على الطبيعة بدقة متناهية.',
    size=10.5, align='justify', space_after=4
)

# ---- Experience ----
main_section_title(main_cell, 'الخبرات العملية')


def experience_entry(cell, role, company, period, bullets):
    # Role | Company line
    p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    ar_run(p, role, size=12, bold=True, color=NAVY)
    ar_run(p, '  ◆  ', size=9, color=ACCENT, bold=True)
    ar_run(p, company, size=12, bold=True, color=ACCENT)

    # Period
    p2 = cell.add_paragraph()
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    ar_run(p2, period, size=9.5, color=GRAY, italic=True)

    for b in bullets:
        main_bullet(cell, b, size=10.5)


experience_entry(
    main_cell,
    'فني مساحة',
    'شركة حسن علام للإنشاءات',
    '2025 - حتى الآن',
    [
        'تنفيذ أعمال التوقيع والرفع المساحي للأعمال الإنشائية وأعمال الطرق والبنية التحتية باستخدام التوتال استيشن وأجهزة الـ GPS / RTK.',
        'إعداد أعمال الرفع الطبوغرافي والـ As-Built والتحقق من مطابقتها للرسومات التنفيذية ومواصفات المشروع.',
        'التنسيق اليومي مع مهندسي الموقع وفريق الجودة والمقاولين من الباطن لضمان سير العمل بدقة وكفاءة عالية.',
        'إعداد تقارير يومية دقيقة لأعمال المساحة وتسليمها لمدير المشروع.',
    ],
)

experience_entry(
    main_cell,
    'مساح موقع',
    'شركة سامكرو (SAMCRO)',
    '2024 - 2025',
    [
        'تنفيذ أعمال التوقيع وضبط النقاط المرجعية للأعمال المدنية والمباني بدقة عالية.',
        'استخدام الميزان والتوتال استيشن في نقل المناسيب وأعمال الحفر والشدات الخشبية.',
        'المساهمة في حساب الكميات والأحجام لأعمال الحفر والردم.',
    ],
)

experience_entry(
    main_cell,
    'مساح',
    'الشركة الوطنية',
    '2023 - 2024',
    [
        'تنفيذ أعمال الرفع الطبوغرافي وإنشاء نقاط تحكم باستخدام أجهزة الـ GPS لمشروعات البنية التحتية.',
        'قراءة الرسومات الإنشائية ونقل الإحداثيات التصميمية إلى الطبيعة بدقة عالية.',
        'المساعدة في إعداد رسومات الـ As-Built وتحديث بيانات الموقع.',
    ],
)

experience_entry(
    main_cell,
    'مساح مبتدئ',
    'شركة رود رنر',
    '2023',
    [
        'المساعدة في أعمال التوقيع والرفع المساحي وجمع البيانات الحقلية في مشروعات الطرق.',
        'تنفيذ قراءات الميزان وأعمال الـ Staking تحت إشراف المساحين الأقدم.',
        'اكتساب أساس قوي في أعمال المساحة بالمواقع والتنسيق بين فرق العمل.',
    ],
)

# ---- Achievements / Highlights ----
main_section_title(main_cell, 'أبرز الإنجازات')
for item in [
    'المشاركة في تنفيذ أعمال المساحة لعدد من المشروعات القومية الكبرى مع شركات المقاولات الرائدة.',
    'تنفيذ أعمال الرفع والتوقيع المساحي بدقة عالية للأعمال الإنشائية والبنية التحتية.',
    'بناء علاقات عمل قوية مع المهندسين والاستشاريين بفضل الالتزام والدقة في الأداء.',
]:
    main_bullet(main_cell, item, size=10.5)

# ---- References ----
main_section_title(main_cell, 'المراجع')
main_paragraph(main_cell, 'تتوفر عند الطلب.', size=10.5, color=GRAY,
               italic=True, align='right', space_after=0)

out = r'e:\apps\restrant\Mohamed_Ali_Salem_CV_AR_PRO.docx'
doc.save(out)
print(f'Professional Arabic CV saved to: {out}')
