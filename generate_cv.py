"""Generate a professional Word CV for Mohamed Ali Salem — Land Surveyor."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Colors & helpers ----------
NAVY = RGBColor(0x0B, 0x2F, 0x5C)      # Deep navy for headings
ACCENT = RGBColor(0x1F, 0x6F, 0xB5)    # Accent blue
DARK = RGBColor(0x1A, 0x1A, 0x1A)      # Body text
GRAY = RGBColor(0x55, 0x55, 0x55)      # Subtle gray

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_horizontal_line(paragraph, color_hex='1F6FB5'):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_horizontal_line(p)
    return p

def body_para(doc, text, bold=False, size=11, color=DARK, space_after=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.runs[0] if p.runs else p.add_run('')
    # Replace with our text
    p.clear() if False else None
    # Simplest: clear existing runs then add
    for r in p.runs:
        r.text = ''
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    return p

# ---------- Document ----------
doc = Document()

# Page margins (tight, clean)
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ---------- Header: Name banner ----------
header_table = doc.add_table(rows=1, cols=1)
header_table.autofit = True
cell = header_table.cell(0, 0)
set_cell_bg(cell, '0B2F5C')

# Clear default paragraph in cell
cell_para = cell.paragraphs[0]
cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
cell_para.paragraph_format.space_before = Pt(6)
cell_para.paragraph_format.space_after = Pt(0)
run = cell_para.add_run('MOHAMED ALI SALEM')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
run.font.name = 'Calibri'

# Subtitle
sub_para = cell.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_before = Pt(0)
sub_para.paragraph_format.space_after = Pt(4)
sub_run = sub_para.add_run('LAND SURVEYOR  |  SITE SURVEYOR')
sub_run.font.size = Pt(12)
sub_run.font.color.rgb = RGBColor(0xCF, 0xE2, 0xF3)
sub_run.font.name = 'Calibri'
sub_run.bold = True

# Contact line
contact_para = cell.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_para.paragraph_format.space_before = Pt(0)
contact_para.paragraph_format.space_after = Pt(6)
contact_run = contact_para.add_run(
    '📞  +20 103 216 9010     ✉  mimo99999911@gmail.com     📍  Cairo, Egypt'
)
contact_run.font.size = Pt(10.5)
contact_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
contact_run.font.name = 'Calibri'

# Small spacer
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------- Professional Summary ----------
section_heading(doc, 'Professional Summary')
body_para(
    doc,
    'Detail-oriented Land Surveyor with over 3 years of hands-on field experience '
    'supporting large-scale construction and infrastructure projects across Egypt. '
    'Skilled in operating Total Station, Auto Level, and GPS / GNSS (RTK) equipment '
    'for accurate setting-out, topographic, and as-built surveys. Trusted contributor '
    'to top-tier contractors including Hassan Allam, Road Runner, Al-Wataniya, and SAMCRO. '
    'Known for precision, reliability on site, and the ability to translate shop drawings '
    'into accurate field execution.'
)

# ---------- Core Skills ----------
section_heading(doc, 'Core Skills & Equipment')

skills_table = doc.add_table(rows=1, cols=2)
skills_table.autofit = True
left_col = skills_table.cell(0, 0)
right_col = skills_table.cell(0, 1)

def skills_cell(cell, items):
    cell.paragraphs[0].text = ''
    for item in items:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run('• ')
        r.font.size = Pt(11)
        r.font.color.rgb = ACCENT
        r.bold = True
        r2 = p.add_run(item)
        r2.font.size = Pt(11)
        r2.font.color.rgb = DARK
        r2.font.name = 'Calibri'

skills_cell(left_col, [
    'Total Station Operation',
    'Auto Level / Digital Level',
    'GPS / GNSS / RTK Surveying',
    'Setting-Out & Layout',
    'Topographic Surveys',
    'As-Built Surveys',
])
skills_cell(right_col, [
    'Quantity & Volume Calculations',
    'Reading Shop & Construction Drawings',
    'Coordinate Systems & Datum Transformation',
    'AutoCAD (2D)',
    'Quality Control on Site',
    'Team Coordination with Engineers',
])

# Remove table borders
tbl = skills_table._tbl
tblPr = tbl.tblPr
tblBorders = OxmlElement('w:tblBorders')
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    b = OxmlElement(f'w:{border_name}')
    b.set(qn('w:val'), 'nil')
    tblBorders.append(b)
tblPr.append(tblBorders)

# ---------- Professional Experience ----------
section_heading(doc, 'Professional Experience')

def experience_block(doc, role, company, period, bullets):
    # Role + Company line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(role)
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = NAVY
    r1.font.name = 'Calibri'
    r2 = p.add_run(f'   |   {company}')
    r2.font.size = Pt(12)
    r2.font.color.rgb = ACCENT
    r2.bold = True
    r2.font.name = 'Calibri'

    # Period line
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    pr = p2.add_run(period)
    pr.italic = True
    pr.font.size = Pt(10.5)
    pr.font.color.rgb = GRAY
    pr.font.name = 'Calibri'

    for b in bullets:
        bullet(doc, b)

experience_block(
    doc,
    'Land Surveyor',
    'Hassan Allam Construction',
    '2025 – Present',
    [
        'Executed precise setting-out for structural, road, and infrastructure works using Total Station and GNSS/RTK GPS.',
        'Produced topographic and as-built surveys, ensuring alignment with project drawings and tolerances.',
        'Coordinated daily with site engineers, QC, and sub-contractors to keep construction on schedule.',
        'Maintained accurate field notes, survey logs, and reported progress to the project surveyor/manager.',
    ]
)

experience_block(
    doc,
    'Site Surveyor',
    'SAMCRO',
    '2024 – 2025',
    [
        'Performed layout and control surveys for building and civil works, delivering accurate benchmarks and reference points.',
        'Operated Auto Level and Total Station for level transfer, excavation depths, and formwork verification.',
        'Supported quantity take-offs and volume calculations for earthworks and backfilling.',
    ]
)

experience_block(
    doc,
    'Surveyor',
    'Al-Wataniya',
    '2023 – 2024',
    [
        'Carried out topographic surveys and GPS-based control point establishment on infrastructure projects.',
        'Reviewed construction drawings and transferred design coordinates to the field with high accuracy.',
        'Assisted in preparing as-built drawings and updating site data.',
    ]
)

experience_block(
    doc,
    'Junior Surveyor',
    'Road Runner',
    '2023',
    [
        'Supported senior surveyors in setting-out works and field data collection on roadway projects.',
        'Handled Auto Level readings, staking, and basic GPS operations under supervision.',
        'Gained strong foundation in site surveying workflows and construction coordination.',
    ]
)

# ---------- Education ----------
section_heading(doc, 'Education')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run('Diploma in Surveying')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = NAVY
r.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run('Graduated: 2023')
r2.italic = True
r2.font.size = Pt(11)
r2.font.color.rgb = GRAY
r2.font.name = 'Calibri'

# ---------- Languages ----------
section_heading(doc, 'Languages')
lang_table = doc.add_table(rows=1, cols=2)
lang_table.autofit = True

def lang_cell(cell, lang, level):
    cell.paragraphs[0].text = ''
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f'{lang}: ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY
    r1.font.name = 'Calibri'
    r2 = p.add_run(level)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    r2.font.name = 'Calibri'

lang_cell(lang_table.cell(0, 0), 'Arabic', 'Native')
lang_cell(lang_table.cell(0, 1), 'English', 'Working proficiency (technical)')

tbl2 = lang_table._tbl
tblPr2 = tbl2.tblPr
tblBorders2 = OxmlElement('w:tblBorders')
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    b = OxmlElement(f'w:{border_name}')
    b.set(qn('w:val'), 'nil')
    tblBorders2.append(b)
tblPr2.append(tblBorders2)

# ---------- Key Strengths ----------
section_heading(doc, 'Key Strengths')
for item in [
    'Precision and attention to detail on every measurement.',
    'Strong field experience with top Egyptian contractors.',
    'Reliable, punctual, and adaptable to demanding project timelines.',
    'Effective communicator with engineers, foremen, and consultants.',
]:
    bullet(doc, item)

# ---------- References ----------
section_heading(doc, 'References')
body_para(doc, 'Available upon request.', color=GRAY, size=11)

# Save
output_path = r'e:\apps\restrant\Mohamed_Ali_Salem_CV.docx'
doc.save(output_path)
print(f'CV saved to: {output_path}')
