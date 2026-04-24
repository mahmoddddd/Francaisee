"""Generate a professional Arabic Word CV for Mohamed Ali Salem — Land Surveyor."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Colors ----------
NAVY = RGBColor(0x0B, 0x2F, 0x5C)
ACCENT = RGBColor(0x1F, 0x6F, 0xB5)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xCF, 0xE2, 0xF3)

AR_FONT = 'Arial'  # Clean, renders Arabic well in Word


# ---------- Helpers ----------
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_rtl(paragraph):
    """Mark paragraph as right-to-left."""
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    p_pr.append(bidi)


def set_run_rtl(run):
    """Mark run as RTL (so Arabic font sizes apply correctly)."""
    r_pr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    r_pr.append(rtl)


def set_run_complex_font(run, font_name, size_pt):
    """Set complex-script font (for Arabic rendering)."""
    r_pr = run._r.get_or_add_rPr()
    rFonts = r_pr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r_pr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    # Complex-script size
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    r_pr.append(szCs)


def ar_run(paragraph, text, size=11, bold=False, color=DARK, font=AR_FONT):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    set_run_complex_font(run, font, size)
    set_run_rtl(run)
    return run


def add_horizontal_line(paragraph, color_hex='1F6FB5'):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def section_heading_ar(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    ar_run(p, text, size=14, bold=True, color=NAVY)
    add_horizontal_line(p)
    return p


def ar_paragraph(doc, text, size=11, bold=False, color=DARK, align='right',
                 space_after=3, space_before=0):
    p = doc.add_paragraph()
    set_rtl(p)
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    ar_run(p, text, size=size, bold=bold, color=color)
    return p


def ar_bullet(doc, text, size=11):
    """Bullet with Arabic-friendly right-aligned marker."""
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    # Bullet first (on the right visually because RTL)
    bullet_run = p.add_run('◀ ')
    bullet_run.font.name = AR_FONT
    bullet_run.font.size = Pt(size)
    bullet_run.font.color.rgb = ACCENT
    bullet_run.bold = True
    set_run_complex_font(bullet_run, AR_FONT, size)
    set_run_rtl(bullet_run)
    ar_run(p, text, size=size, color=DARK)
    return p


# ---------- Document ----------
doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Set document default to RTL
# (Applied per paragraph via set_rtl)

# Default style font
style = doc.styles['Normal']
style.font.name = AR_FONT
style.font.size = Pt(11)

# ---------- Header Banner ----------
header_table = doc.add_table(rows=1, cols=1)
header_table.autofit = True
cell = header_table.cell(0, 0)
set_cell_bg(cell, '0B2F5C')

# Name
p_name = cell.paragraphs[0]
set_rtl(p_name)
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_name.paragraph_format.space_before = Pt(6)
p_name.paragraph_format.space_after = Pt(0)
ar_run(p_name, 'محمد علي سالم', size=28, bold=True, color=WHITE)

# Title
p_title = cell.add_paragraph()
set_rtl(p_title)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(4)
ar_run(p_title, 'مهندس / فني مساحة', size=14, bold=True, color=LIGHT_BLUE)

# Contact
p_contact = cell.add_paragraph()
set_rtl(p_contact)
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contact.paragraph_format.space_before = Pt(0)
p_contact.paragraph_format.space_after = Pt(6)
ar_run(
    p_contact,
    '📞  01032169010     ✉  mimo99999911@gmail.com     📍  القاهرة - مصر',
    size=11, color=WHITE
)

# Spacer
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(2)

# ---------- Professional Summary ----------
section_heading_ar(doc, 'الملخص المهني')
ar_paragraph(
    doc,
    'فني مساحة صاحب خبرة عملية تزيد عن 3 سنوات في تنفيذ أعمال المساحة لمشروعات '
    'إنشائية وبنية تحتية كبرى داخل جمهورية مصر العربية. يتقن تشغيل أجهزة الميزان '
    '(Auto Level) والتوتال استيشن (Total Station) وأجهزة الـ GPS / GNSS (RTK) '
    'في أعمال الرفع المساحي، ونقل المناسيب، وتوقيع النقاط، وأعمال As-Built. '
    'عمل مع كبرى شركات المقاولات في مصر مثل حسن علام، رود رنر، الوطنية، وسامكرو. '
    'يتميز بالدقة في العمل، والالتزام بالمواعيد، والقدرة على قراءة الرسومات التنفيذية '
    'وتنفيذها على الطبيعة بدقة عالية.',
    size=11, align='justify', space_after=4
)

# ---------- Skills ----------
section_heading_ar(doc, 'المهارات والأجهزة')

skills_left = [
    'تشغيل جهاز التوتال استيشن (Total Station)',
    'تشغيل جهاز الميزان (Auto Level / Digital Level)',
    'تشغيل أجهزة الـ GPS / GNSS / RTK',
    'أعمال التوقيع والرفع المساحي (Setting-Out)',
    'الرفع الطبوغرافي (Topographic Surveys)',
    'أعمال الـ As-Built للمشروعات المنفذة',
]
skills_right = [
    'حساب الكميات والمناسيب والحفر والردم',
    'قراءة الرسومات الإنشائية والتنفيذية (Shop Drawings)',
    'التعامل مع أنظمة الإحداثيات وتحويلها',
    'برنامج AutoCAD (ثنائي الأبعاد 2D)',
    'أعمال الجودة والتحقق من الدقة على الموقع',
    'التنسيق مع المهندسين والاستشاريين والفرق الفنية',
]

skills_table = doc.add_table(rows=1, cols=2)
skills_table.autofit = True
left_cell = skills_table.cell(0, 0)
right_cell = skills_table.cell(0, 1)

def fill_skills_cell(cell, items):
    cell.paragraphs[0].text = ''
    for item in items:
        p = cell.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(1)
        bullet_r = p.add_run('◀ ')
        bullet_r.font.name = AR_FONT
        bullet_r.font.size = Pt(11)
        bullet_r.font.color.rgb = ACCENT
        bullet_r.bold = True
        set_run_complex_font(bullet_r, AR_FONT, 11)
        set_run_rtl(bullet_r)
        ar_run(p, item, size=11, color=DARK)

# In RTL, the "right" column visually appears on the right. But since python-docx
# tables aren't RTL-aware, fill both so layout reads nicely.
fill_skills_cell(right_cell, skills_right)
fill_skills_cell(left_cell, skills_left)

# Remove borders
tbl = skills_table._tbl
tblPr = tbl.tblPr
tblBorders = OxmlElement('w:tblBorders')
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    b = OxmlElement(f'w:{border_name}')
    b.set(qn('w:val'), 'nil')
    tblBorders.append(b)
tblPr.append(tblBorders)

# Make table RTL
tbl_bidi = OxmlElement('w:bidiVisual')
tbl_bidi.set(qn('w:val'), '1')
tblPr.append(tbl_bidi)

# ---------- Experience ----------
section_heading_ar(doc, 'الخبرات العملية')


def experience_block_ar(doc, role, company, period, bullets):
    # Role + Company
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    ar_run(p, role, size=13, bold=True, color=NAVY)
    ar_run(p, '   |   ', size=12, color=GRAY)
    ar_run(p, company, size=13, bold=True, color=ACCENT)

    # Period
    p2 = doc.add_paragraph()
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r = ar_run(p2, period, size=11, color=GRAY)
    r.italic = True

    for b in bullets:
        ar_bullet(doc, b)


experience_block_ar(
    doc,
    'فني مساحة',
    'شركة حسن علام للإنشاءات',
    '2025 - حتى الآن',
    [
        'تنفيذ أعمال التوقيع والرفع المساحي للأعمال الإنشائية والطرق والبنية التحتية باستخدام التوتال استيشن وأجهزة الـ GPS / RTK.',
        'إعداد أعمال الرفع الطبوغرافي والـ As-Built والتأكد من مطابقتها للرسومات التنفيذية.',
        'التنسيق اليومي مع مهندسي الموقع وقسم الجودة والمقاولين من الباطن لضمان سير العمل بدقة.',
        'الاحتفاظ بسجلات دقيقة لأعمال المساحة وتقديم التقارير اليومية لمشرف المشروع.',
    ],
)

experience_block_ar(
    doc,
    'مساح موقع',
    'شركة سامكرو (SAMCRO)',
    '2024 - 2025',
    [
        'تنفيذ أعمال التوقيع وضبط النقاط المرجعية للأعمال المدنية والمباني بدقة عالية.',
        'استخدام الميزان والتوتال استيشن في نقل المناسيب وأعمال الحفر والشدات الخشبية.',
        'المساهمة في حساب الكميات والأحجام لأعمال الحفر والردم.',
    ],
)

experience_block_ar(
    doc,
    'مساح',
    'الشركة الوطنية',
    '2023 - 2024',
    [
        'تنفيذ أعمال الرفع الطبوغرافي وإنشاء نقاط تحكم باستخدام أجهزة الـ GPS لمشروعات البنية التحتية.',
        'قراءة الرسومات الإنشائية ونقل الإحداثيات التصميمية إلى الطبيعة بدقة عالية.',
        'المساعدة في إعداد رسومات الـ As-Built وتحديث بيانات الموقع.',
    ],
)

experience_block_ar(
    doc,
    'مساح مبتدئ',
    'شركة رود رنر',
    '2023',
    [
        'المساعدة في أعمال التوقيع والرفع المساحي وجمع البيانات الحقلية في مشروعات الطرق.',
        'تنفيذ قراءات الميزان وأعمال الـ Staking الأساسية تحت إشراف المساحين الأقدم.',
        'اكتساب أساس قوي في أعمال المساحة بالمواقع والتنسيق بين فرق العمل.',
    ],
)

# ---------- Education ----------
section_heading_ar(doc, 'المؤهلات العلمية')
p = doc.add_paragraph()
set_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(0)
ar_run(p, 'دبلوم مساحة', size=13, bold=True, color=NAVY)

p2 = doc.add_paragraph()
set_rtl(p2)
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.paragraph_format.space_after = Pt(2)
r = ar_run(p2, 'سنة التخرج: 2023', size=11, color=GRAY)
r.italic = True

# ---------- Languages ----------
section_heading_ar(doc, 'اللغات')
lang_table = doc.add_table(rows=1, cols=2)
lang_table.autofit = True


def lang_cell_ar(cell, lang, level):
    cell.paragraphs[0].text = ''
    p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    ar_run(p, f'{lang}: ', size=12, bold=True, color=NAVY)
    ar_run(p, level, size=12, color=DARK)


lang_cell_ar(lang_table.cell(0, 1), 'اللغة العربية', 'اللغة الأم')
lang_cell_ar(lang_table.cell(0, 0), 'اللغة الإنجليزية', 'مستوى جيد (فني / تقني)')

tbl2 = lang_table._tbl
tblPr2 = tbl2.tblPr
tblBorders2 = OxmlElement('w:tblBorders')
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    b = OxmlElement(f'w:{border_name}')
    b.set(qn('w:val'), 'nil')
    tblBorders2.append(b)
tblPr2.append(tblBorders2)
tbl2_bidi = OxmlElement('w:bidiVisual')
tbl2_bidi.set(qn('w:val'), '1')
tblPr2.append(tbl2_bidi)

# ---------- Strengths ----------
section_heading_ar(doc, 'نقاط القوة')
for item in [
    'الدقة العالية والاهتمام بأدق التفاصيل في كل القياسات.',
    'خبرة ميدانية قوية مع كبرى شركات المقاولات في مصر.',
    'الالتزام وسرعة الاستجابة تحت ضغط المواعيد النهائية للمشروعات.',
    'مهارات تواصل جيدة مع المهندسين والمشرفين والاستشاريين.',
    'قدرة عالية على العمل ضمن فريق وإدارة الأولويات في الموقع.',
]:
    ar_bullet(doc, item)

# ---------- References ----------
section_heading_ar(doc, 'الإشارات المرجعية')
ar_paragraph(doc, 'تتوفر عند الطلب.', size=11, color=GRAY)

# Save
out = r'e:\apps\restrant\Mohamed_Ali_Salem_CV_AR.docx'
doc.save(out)
print(f'Arabic CV saved to: {out}')
